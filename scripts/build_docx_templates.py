"""Build seed docx templates with docxtpl-compatible placeholders.

These templates are intentionally simple. They give each department/document
type a separate physical template file now, and can be replaced by formal
company templates later without changing skill code.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_ROOT = ROOT / "assets" / "templates" / "docx"


TEMPLATES = {
    "admin": {
        "notice": {"label": "通知", "title_size": 22, "font": "SimHei"},
        "opinion": {"label": "意见", "title_size": 22, "font": "SimHei"},
        "action_plan": {"label": "方案 / 行动计划", "title_size": 22, "font": "SimHei"},
        "approval": {"label": "批复", "title_size": 22, "font": "SimHei"},
    },
    "marketing": {
        "action_plan": {"label": "营销方案", "title_size": 20, "font": "Microsoft YaHei"},
        "opinion": {"label": "营销意见", "title_size": 20, "font": "Microsoft YaHei"},
    },
    "product": {
        "action_plan": {"label": "产品方案", "title_size": 20, "font": "Microsoft YaHei"},
    },
}


def set_run_font(run, font_name: str, size: int | None = None, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def configure_document(doc: Document, body_font: str = "SimSun") -> None:
    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    style = doc.styles["Normal"]
    style.font.name = body_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)


def set_body_paragraph(paragraph, indent: bool = True) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if indent:
        paragraph.paragraph_format.first_line_indent = Pt(28)


def add_template_body(doc: Document, config: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    title_run = title.add_run("{{ title }}")
    set_run_font(title_run, config["font"], size=config["title_size"], bold=True)

    recipient = doc.add_paragraph("{{ recipient }}：")
    set_body_paragraph(recipient, indent=False)

    doc.add_paragraph("{%p for paragraph in paragraphs %}")
    body = doc.add_paragraph("{{ paragraph }}")
    set_body_paragraph(body)
    doc.add_paragraph("{%p endfor %}")

    doc.add_paragraph("{%p for section in sections %}")
    section_heading = doc.add_paragraph("{{ section.heading }}")
    section_heading.paragraph_format.line_spacing = 1.5
    section_heading.paragraph_format.space_before = Pt(6)
    section_heading.paragraph_format.space_after = Pt(0)
    set_run_font(section_heading.runs[0], config["font"], size=14, bold=True)

    doc.add_paragraph("{%p for item in section.lines %}")
    item = doc.add_paragraph("{{ item }}")
    set_body_paragraph(item)
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endfor %}")

    closing = doc.add_paragraph("{{ closing }}")
    set_body_paragraph(closing)
    doc.add_paragraph()

    signer = doc.add_paragraph("{{ signer }}")
    signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signer.paragraph_format.line_spacing = 1.5
    date = doc.add_paragraph("{{ date }}")
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.paragraph_format.line_spacing = 1.5


def build() -> None:
    for department, templates in TEMPLATES.items():
        department_dir = TEMPLATE_ROOT / department
        department_dir.mkdir(parents=True, exist_ok=True)
        for document_type, config in templates.items():
            doc = Document()
            configure_document(doc)
            add_template_body(doc, config)
            doc.save(department_dir / f"{document_type}.docx")


if __name__ == "__main__":
    build()
