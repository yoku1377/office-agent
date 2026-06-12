"""公文润色 skill 主流程。

用法：
  python -m skills.polish.polish 输入.docx --level light|medium|heavy \
      [--terms assets/terms/terms.yaml] [--out 输出.docx]

离线演示（无需 API Key）：
  MOCK_LLM_FILE=tests/sample/mock_response.json python -m skills.polish.polish ...
"""
import argparse
import json
import os
import re
import sys

import yaml
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))
from app.providers.llm import get_llm  # noqa: E402
from skills.polish.prompts import build_prompt  # noqa: E402
from skills.polish.writeback import apply_revisions  # noqa: E402


def load_terms(path):
    if not path or not os.path.exists(path):
        return []
    data = yaml.safe_load(open(path, encoding="utf-8")) or {}
    return list(data.get("terms", []))


def extract_paragraphs(doc):
    """返回 [(index, text)]，跳过空段。index 与 doc.paragraphs 对齐。"""
    return [(i, p.text) for i, p in enumerate(doc.paragraphs) if p.text.strip()]


def parse_llm_json(raw):
    raw = re.sub(r"^```(json)?|```$", "", raw.strip(), flags=re.M).strip()
    return json.loads(raw)


def validate(revisions, paragraphs, terms):
    """三道闸：段落存在、原文可定位、术语未被改动。返回 (通过, 拒绝及原因)。"""
    text_by_index = dict(paragraphs)
    ok, rejected = [], []
    for rev in revisions:
        i = int(rev.get("para_index", -1))
        if i not in text_by_index:
            rejected.append((rev, "段落索引不存在"))
            continue
        if rev["old"] not in text_by_index[i]:
            rejected.append((rev, "原文在段落中未找到"))
            continue
        broken = [t for t in terms if t in rev["old"] and t not in rev["new"]]
        if broken:
            rejected.append((rev, f"术语被改动: {broken}"))
            continue
        ok.append(rev)
    return ok, rejected


def _fix_settings(doc):
    """python-docx 默认模板的 w:zoom 缺少 percent 属性，补齐以通过 OOXML 严格校验。"""
    from docx.oxml.ns import qn
    settings = doc.settings.element
    for zoom in settings.findall(qn("w:zoom")):
        if zoom.get(qn("w:percent")) is None:
            zoom.set(qn("w:percent"), "100")


def polish(
    in_path,
    level="medium",
    terms_path=None,
    terms=None,
    style_card=None,
    out_path=None,
    author="AI润色",
):
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    loaded_terms = load_terms(terms_path)
    terms = list(dict.fromkeys(loaded_terms + list(terms or [])))

    system, user = build_prompt(paragraphs, level, terms, style_card=style_card)
    raw = get_llm().chat(system, user)
    revisions = parse_llm_json(raw)

    ok, rejected = validate(revisions, paragraphs, terms)
    applied = apply_revisions(doc, ok, author=author)

    _fix_settings(doc)
    out_path = out_path or in_path.replace(".docx", f"_润色_{level}.docx")
    doc.save(out_path)
    return {"out": out_path, "applied": applied, "rejected": [(r["old"], why) for r, why in rejected]}


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("--level", default="medium", choices=["light", "medium", "heavy"])
    ap.add_argument("--terms", default="assets/terms/terms.yaml")
    ap.add_argument("--out", default=None)
    args = ap.parse_args()
    result = polish(args.input, level=args.level, terms_path=args.terms, out_path=args.out)
    print(json.dumps(result, ensure_ascii=False, indent=2))
