"""Render structured document JSON into a Word file."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def build_template_context(data: dict) -> dict:
    sections = []
    for section in data.get("sections") or []:
        sections.append(
            {
                "heading": str(section.get("heading") or "").strip(),
                "lines": [str(item).strip() for item in section.get("items") or [] if str(item).strip()],
            }
        )
    return {
        "title": data.get("title") or "待补充",
        "recipient": data.get("recipient") or "待补充",
        "paragraphs": [str(text).strip() for text in data.get("paragraphs") or [] if str(text).strip()],
        "sections": sections,
        "closing": data.get("closing") or "",
        "signer": data.get("signer") or "待补充",
        "date": data.get("date") or "待补充",
    }


def render_template_docx(data: dict, template_path: str, out_path: str) -> str:
    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError("docxtpl is not installed") from exc

    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(template_path)

    output = Path(out_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxTemplate(str(template))
    doc.render(build_template_context(data))
    doc.save(output)
    return str(output)


def _set_run_font(run, font_name: str, size: int | None = None, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)


def _add_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(28)


def render_plain_docx(data: dict, out_path: str) -> str:
    doc = Document()
    _configure_document(doc)

    title = data.get("title") or "待补充"
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(18)
    title_run = title_p.add_run(title)
    _set_run_font(title_run, "SimHei", size=22, bold=True)

    recipient = data.get("recipient")
    if recipient:
        _add_paragraph(doc, f"{recipient}：", first_line_indent=False)

    for text in data.get("paragraphs") or []:
        if str(text).strip():
            _add_paragraph(doc, str(text).strip())

    for section in data.get("sections") or []:
        heading = str(section.get("heading") or "").strip()
        if heading:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(heading)
            _set_run_font(run, "SimHei", size=14, bold=True)
        for item in section.get("items") or []:
            item_text = str(item).strip()
            if item_text:
                _add_paragraph(doc, item_text)

    closing = data.get("closing")
    if closing:
        _add_paragraph(doc, str(closing).strip())

    signer = data.get("signer")
    date = data.get("date")
    if signer or date:
        doc.add_paragraph()
        if signer:
            p = doc.add_paragraph(str(signer).strip())
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if date:
            p = doc.add_paragraph(str(date).strip())
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return str(path)


def render_docx(data: dict, out_path: str, template_path: str | None = None) -> dict:
    if template_path:
        try:
            output = render_template_docx(data, template_path, out_path)
            return {"out": output, "engine": "docxtpl", "template_path": template_path}
        except Exception as exc:
            output = render_plain_docx(data, out_path)
            return {
                "out": output,
                "engine": "python-docx",
                "template_path": template_path,
                "template_error": str(exc),
            }

    output = render_plain_docx(data, out_path)
    return {"out": output, "engine": "python-docx", "template_path": None}
