"""生成带有典型问题的样例公文，用于离线验证润色流水线。"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "SimSun"
style.font.size = Pt(14)

t = doc.add_paragraph("关于开展红途智盒 IB-200 现场巡检工作的通知")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.bold = True

doc.add_paragraph("各部门：")
doc.add_paragraph("为了进一步的加强设备运行的安全性的管理，根据公司的相关的规定，现就开展红途智盒 IB-200 现场巡检工作有关事项通知如下。")
doc.add_paragraph("一、巡检范围。本次巡检覆盖所有已部署红途智盒 IB-200 的现场，包括在Ascend平台上运行的全部业务系统，确保设备和系统都能够正常的运行没有问题。")
doc.add_paragraph("二、时间按排。巡检工作自6月15日开始，至6月30日结束，请各部门提前做好相关准备工作的安排。")
doc.add_paragraph("三、工作要求。各部门要高度重视，认真组织，确保巡检工作顺利完成，发现问题要及时上报并积极的进行整改。")
doc.add_paragraph("特此通知。")
doc.save("tests/sample/sample_notice.docx")
print("sample written")
